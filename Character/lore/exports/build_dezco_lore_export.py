#!/usr/bin/env python3
"""Build downloadable Dezco lore exports (tabbed spreadsheet + Google Docs ready docx)."""

from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.text import WD_BREAK
from docx.shared import Inches, Pt, RGBColor
from openpyxl import Workbook
from openpyxl.styles import Alignment, Font, PatternFill
from openpyxl.utils import get_column_letter

ROOT = Path(__file__).resolve().parents[3]
LORE = ROOT / "character" / "lore"
CRESTS = ROOT / "assets" / "crests" / "alzenhiem"
OUT = Path(__file__).resolve().parent
ARTIFACTS = Path("/opt/cursor/artifacts/dezco-lore-export")

HEADER_FILL = PatternFill("solid", fgColor="1F2937")
HEADER_FONT = Font(color="FFFFFF", bold=True)
TITLE_FONT = Font(bold=True, size=14, color="111827")
SECTION_FONT = Font(bold=True, size=12, color="1F2937")
WRAP = Alignment(wrap_text=True, vertical="top")


def read(path: Path) -> str:
    return path.read_text(encoding="utf-8").strip() + "\n"


def style_header(ws, cols: int):
    for c in range(1, cols + 1):
        cell = ws.cell(1, c)
        cell.fill = HEADER_FILL
        cell.font = HEADER_FONT
        cell.alignment = Alignment(wrap_text=True, vertical="center")


def autosize(ws, min_w=12, max_w=60):
    for col in ws.columns:
        letter = get_column_letter(col[0].column)
        length = 0
        for cell in col:
            if cell.value is None:
                continue
            length = max(length, min(max_w, max(len(str(cell.value).split("\n")[0]), min_w)))
        ws.column_dimensions[letter].width = length + 2


def add_text_sheet(wb: Workbook, title: str, text: str, source: str = ""):
    ws = wb.create_sheet(title[:31])
    ws["A1"] = title
    ws["A1"].font = TITLE_FONT
    ws["A2"] = f"Source: {source}" if source else ""
    ws["A2"].font = Font(italic=True, color="6B7280")
    row = 4
    for line in text.splitlines():
        ws.cell(row, 1, line)
        ws.cell(row, 1).alignment = WRAP
        row += 1
    ws.column_dimensions["A"].width = 100
    ws.row_dimensions[1].height = 22


def add_table_sheet(wb: Workbook, title: str, headers: list[str], rows: list[list[str]], note: str = ""):
    ws = wb.create_sheet(title[:31])
    ws["A1"] = title
    ws["A1"].font = TITLE_FONT
    start = 3
    if note:
        ws["A2"] = note
        ws["A2"].font = Font(italic=True, color="6B7280")
        start = 3
    for i, h in enumerate(headers, 1):
        ws.cell(start, i, h)
    style_header(ws if start == 1 else ws, len(headers))
    # restyle correct header row
    for i, h in enumerate(headers, 1):
        cell = ws.cell(start, i, h)
        cell.fill = HEADER_FILL
        cell.font = HEADER_FONT
        cell.alignment = WRAP
    for r_i, row in enumerate(rows, start + 1):
        for c_i, val in enumerate(row, 1):
            cell = ws.cell(r_i, c_i, val)
            cell.alignment = WRAP
    autosize(ws)
    return ws


def build_workbook() -> Path:
    wb = Workbook()
    cover = wb.active
    cover.title = "00 Index"
    cover["A1"] = "Dezco Knoleburn Lore Pack"
    cover["A1"].font = TITLE_FONT
    cover["A2"] = "Upload this .xlsx to Google Drive and open with Google Sheets for tabbed browsing."
    cover["A3"] = "Long prose also lives in Dezco_Knoleburn_Lore.docx for Google Docs."
    cover["A5"] = "Tabs"
    cover["A5"].font = SECTION_FONT
    tabs = [
        ("00 Index", "This cover and map"),
        ("01 Party Quick Ref", "Short memory sheet for party talk"),
        ("02 Me At A Glance", "Identity stats"),
        ("03 Family", "Parents and siblings"),
        ("04 Names", "Name meanings"),
        ("05 Timeline", "Life events in order"),
        ("06 Places And Gear", "Locations and inventory"),
        ("07 Crests And Sigils", "Asset catalog + meanings"),
        ("08 Abilities Mana Well", "Yua combat magic system"),
        ("09 Secrets", "What to share / not spoil"),
        ("10 Character Intro", "Full character introduction"),
        ("11 Journal 01", "After the Charge"),
        ("12 Journal 02", "Leaving Home for Saltrock"),
        ("13 Journal 03", "The Staff Falls"),
        ("14 Yua Red And Black", "Later color reveal"),
        ("15 Yua Mana Well", "Full Mana Well plot"),
        ("16 Crest Lore Full", "Full crest/sigil writeups"),
    ]
    cover["A6"] = "Tab"
    cover["B6"] = "Contents"
    style_header(cover, 2)
    cover["A6"].fill = HEADER_FILL
    cover["B6"].fill = HEADER_FILL
    cover["A6"].font = HEADER_FONT
    cover["B6"].font = HEADER_FONT
    for i, (name, desc) in enumerate(tabs, 7):
        cover.cell(i, 1, name)
        cover.cell(i, 2, desc)
    cover.column_dimensions["A"].width = 28
    cover.column_dimensions["B"].width = 48

    # Party quick ref as whole text
    add_text_sheet(wb, "01 Party Quick Ref", read(LORE / "dezco-party-reference.md"), "character/lore/dezco-party-reference.md")

    add_table_sheet(
        wb,
        "02 Me At A Glance",
        ["Field", "Value"],
        [
            ["Full name", "Dezco Knoleburn"],
            ["Name means", "stands firm; one who holds his ground"],
            ["Age", "17"],
            ["Birthday", "November 6"],
            ["Height", "5'10\""],
            ["Build", "slender"],
            ["Hair", "black"],
            ["Favorite colors", "red and black"],
            ["Ancestry", "Dwarven + Elven"],
            ["Home", "Alzenhiem (hamlet)"],
            ["Title now", "Hero (given by Yua at 17)"],
            ["Master / trainer", "Rilock (royal master swordsman; King’s Lion royal guard)"],
            ["Directed by", "The King’s Lion (ordered Rilock to train me)"],
            ["Companion", "Yua (often appears as a small pixie)"],
            ["Crest colors", "crimson red + deep black"],
            ["Sigil", "vertical blade / flame split red and black"],
            ["Current location", "Saltrock general building (doors sealed; under attack)"],
            ["Travel status", "Adventure day 2 gone wrong; protective staff fell; mid siege"],
        ],
        "Quick identity card for party talk.",
    )

    add_table_sheet(
        wb,
        "03 Family",
        ["Person", "Who", "Age", "Name means", "Said", "Colors", "One line"],
        [
            ["Dolkin", "Father", "38", "steadfast of kin", "Dol kin", "charcoal, iron, bronze", "Steady, Dwarven depth, sent Dezco to train to keep him alive"],
            ["Fulfein", "Mother", "36", "full fair light", "Fule fe in", "sage, silver, pale gold", "Warm Elven grace; Alzenhiem’s center for Dezco"],
            ["Kalhien", "Younger brother", "15", "rises in pursuit", "Kal he en", "ink blue + copper", "Restless, proud, does not want to be only Dezco’s brother"],
            ["Astiale", "Younger sister", "8", "gentle star", "as T ale", "rose, cream, sky", "Sweet; not born when Dezco left at 8"],
        ],
        "Knoleburn = stream by the knoll. Alzenhiem = sheltered home under a kind sky.",
    )

    add_table_sheet(
        wb,
        "04 Names",
        ["Name", "Said", "Meaning", "Why given"],
        [
            ["Alzenhiem", "Al zen heim", "sheltered home under a kind sky", "Promise of roof and welcome before greatness"],
            ["Knoleburn", "Knole burn", "stream by the knoll", "House tied to living land"],
            ["Dolkin", "Dol kin", "steadfast of kin", "Strength meant to be spent on family and duty"],
            ["Fulfein", "Fule fe in", "full fair light", "Warmth that fills a room without demanding it"],
            ["Dezco", "Dezco", "stands firm", "A name that can survive hard weather"],
            ["Kalhien", "Kal he en", "rises in pursuit", "Honor his hunger to climb his own path"],
            ["Astiale", "as T ale", "gentle star", "Shine kindly; rise without hardening"],
        ],
        "Full writeups: character/lore/knoleburn-names.md",
    )

    add_table_sheet(
        wb,
        "05 Timeline",
        ["Age / when", "What happened"],
        [
            ["0 / Nov 6", "Born; named Dezco; red/black crest + sigil"],
            ["8", "Left Alzenhiem for military training; Kalhien was 6; Astiale not born yet"],
            ["8 to 12", "Overshadowed by rival; Yua watched unseen"],
            ["12", "Rival duel; Yua secretly blessed speed/strength; King’s Lion saw her"],
            ["12", "King’s Lion directed Rilock (royal master swordsman / royal guard) to train Dezco"],
            ["12 to 17", "Real missions under Rilock; false Hero of centuries rumors"],
            ["Before 15", "Name known across much of the realm"],
            ["17", "Master Rilock said training was complete"],
            ["17", "Yua gave ancient title Hero + sacred mission + Mana Well (unexplained); Lion and Rilock present"],
            ["17", "Final goodbye at Alzenhiem; family gifts"],
            ["17", "Reached Saltrock; tavern lodging"],
            ["17", "Protective staff fell; monster attack"],
            ["17", "Accidental mana dash / faceplant; elf girl archer save"],
            ["Now", "Sealed in Saltrock general building; journal 03 open ended"],
        ],
    )

    add_table_sheet(
        wb,
        "06 Places And Gear",
        ["Type", "Name", "Details"],
        [
            ["Place", "Alzenhiem", "Home hamlet; Dezco’s center"],
            ["Place", "Saltrock", "Traveling town; many come and go"],
            ["Place", "Saltrock general building", "Protective staff location; currently sealed inside during attack"],
            ["Place", "Saltrock protective staff", "Used to float and block named monsters; has fallen"],
            ["Place", "Saltrock tavern", "Lodged night before attack; nostalgia bread"],
            ["Gear", "New sword", "From Dolkin; hand forged goodbye"],
            ["Gear", "First smithing tools", "From Dolkin; master of your steel"],
            ["Gear", "10 gold pieces", "From Fulfein"],
            ["Gear", "Cooking pot", "From Fulfein"],
            ["Gear", "New clothing", "From Fulfein; worn out of respect"],
            ["Gear", "Shovel", "From Fulfein; funny roadside duty + darker second meaning"],
        ],
    )

    add_table_sheet(
        wb,
        "07 Crests And Sigils",
        ["Who", "Kind", "File", "Colors", "Motif / meaning", "GitHub"],
        [
            ["Alzenhiem", "Town flag", "alzenhiem-town-crest.png", "gold, green, brown, cream", "Roof, wheat, lane; sheltered home", "https://github.com/Dralorex/General-Questions/blob/main/assets/crests/alzenhiem/alzenhiem-town-crest.png"],
            ["Dolkin", "Crest", "dolkin-crest.png", "charcoal, iron, bronze", "Stone arch + hammer; bear weight", "https://github.com/Dralorex/General-Questions/blob/main/assets/crests/alzenhiem/dolkin-crest.png"],
            ["Dolkin", "Sigil", "dolkin-sigil.png", "charcoal, bronze", "Hammer in arch", "https://github.com/Dralorex/General-Questions/blob/main/assets/crests/alzenhiem/dolkin-sigil.png"],
            ["Fulfein", "Crest", "fulfein-crest.png", "sage, silver, pale gold", "Leaf + crescent; care and watch", "https://github.com/Dralorex/General-Questions/blob/main/assets/crests/alzenhiem/fulfein-crest.png"],
            ["Fulfein", "Sigil", "fulfein-sigil.png", "sage, silver", "Leaf into crescent", "https://github.com/Dralorex/General-Questions/blob/main/assets/crests/alzenhiem/fulfein-sigil.png"],
            ["Dezco", "Crest", "dezco-crest.png", "red, black", "Balanced blade/flame between parent marks", "https://github.com/Dralorex/General-Questions/blob/main/assets/crests/alzenhiem/dezco-crest.png"],
            ["Dezco", "Sigil", "dezco-sigil.png", "red, black", "Split blade/flame", "https://github.com/Dralorex/General-Questions/blob/main/assets/crests/alzenhiem/dezco-sigil.png"],
            ["Kalhien", "Crest", "kalhien-crest.png", "ink blue, copper", "Forward eager mark; own path", "https://github.com/Dralorex/General-Questions/blob/main/assets/crests/alzenhiem/kalhien-crest.png"],
            ["Kalhien", "Sigil", "kalhien-sigil.png", "blue, copper", "Forward chevron", "https://github.com/Dralorex/General-Questions/blob/main/assets/crests/alzenhiem/kalhien-sigil.png"],
            ["Astiale", "Crest", "astiale-crest.png", "rose, cream, sky", "Soft bud/star; grow kindly", "https://github.com/Dralorex/General-Questions/blob/main/assets/crests/alzenhiem/astiale-crest.png"],
            ["Astiale", "Sigil", "astiale-sigil.png", "rose, cream", "Soft bud/star seal", "https://github.com/Dralorex/General-Questions/blob/main/assets/crests/alzenhiem/astiale-sigil.png"],
        ],
        "Tradition: child crest = name + gender + both parents; simpler; unique. Sigils made same day.",
    )

    add_table_sheet(
        wb,
        "08 Abilities Mana Well",
        ["Piece", "Quick fact"],
        [
            ["Source", "Bestowed by Yua with the Hero charge; unexplained at first"],
            ["Well", "Stored mana pool within soul confines; separate from the soul"],
            ["Simple imbue", "Innate; should feel as easy as breathing once understood"],
            ["Sword effect", "Lighter, sharper; short burst of swiftness / agility"],
            ["First use", "Saltrock battle (journal 03); accidental ~15 ft dash mid swing = faceplant"],
            ["Training needed", "Control timing, avoid waste, move with the burst instead of being thrown"],
            ["Yua’s plan", "Stay quiet until Dezco asks or figures it out; changed after watching the biff"],
        ],
        "Full plot: character/lore/yua-story-plots/yua-mana-well.md",
    )

    add_table_sheet(
        wb,
        "09 Secrets",
        ["Topic", "Party can know?", "Notes"],
        [
            ["Alzenhiem + family basics", "Yes", "Safe to share"],
            ["Trained since 8 under the Lion", "Yes", "Safe to share"],
            ["Early Hero of centuries rumors", "Yes", "Emphasize it was rumor then"],
            ["Yua made him Hero at 17", "Yes, if trust", "Often looks like a pixie"],
            ["Exact red/black blood meaning", "No", "Later Yua reveal"],
            ["Mana Well / sword imbue details", "Not yet / soon", "Yua explains after Saltrock faceplant"],
            ["Birthday November 6", "Yes", "Safe"],
        ],
        "DM truths: Black = Dwarven caves/void/dark. Red = Elven holy blood after ill. Together = balance. Mana Well separate from soul.",
    )

    add_text_sheet(wb, "10 Character Intro", read(LORE / "dezco-knoleburn-character.md"), "character/lore/dezco-knoleburn-character.md")
    add_text_sheet(wb, "11 Journal 01", read(LORE / "journal" / "01-after-the-charge.md"), "character/lore/journal/01-after-the-charge.md")
    add_text_sheet(wb, "12 Journal 02", read(LORE / "journal" / "02-leaving-home-for-saltrock.md"), "character/lore/journal/02-leaving-home-for-saltrock.md")
    add_text_sheet(wb, "13 Journal 03", read(LORE / "journal" / "03-saltrock-staff-falls.md"), "character/lore/journal/03-saltrock-staff-falls.md")
    add_text_sheet(wb, "14 Yua Red And Black", read(LORE / "yua-story-plots" / "yua-on-red-and-black.md"), "character/lore/yua-story-plots/yua-on-red-and-black.md")
    add_text_sheet(wb, "15 Yua Mana Well", read(LORE / "yua-story-plots" / "yua-mana-well.md"), "character/lore/yua-story-plots/yua-mana-well.md")
    add_text_sheet(wb, "16 Crest Lore Full", read(CRESTS / "README.md"), "assets/crests/alzenhiem/README.md")

    out = OUT / "Dezco_Knoleburn_Lore_Tabs.xlsx"
    wb.save(out)
    return out


def add_heading(doc: Document, text: str, level: int = 1):
    doc.add_heading(text, level=level)


def add_md_ish(doc: Document, text: str):
    for line in text.splitlines():
        if line.startswith("# "):
            doc.add_heading(line[2:].strip(), level=1)
        elif line.startswith("## "):
            doc.add_heading(line[3:].strip(), level=2)
        elif line.startswith("### "):
            doc.add_heading(line[4:].strip(), level=3)
        elif line.startswith("|") or line.strip() == "---":
            # keep tables as plain preformatted lines for reliable import
            p = doc.add_paragraph(line)
            for run in p.runs:
                run.font.name = "Courier New"
                run.font.size = Pt(9)
        elif line.startswith("> "):
            p = doc.add_paragraph(line[2:].strip())
            p.style = "Intense Quote" if "Intense Quote" in [s.name for s in doc.styles] else p.style
        else:
            doc.add_paragraph(line)


def build_docx() -> Path:
    doc = Document()
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(11)

    add_heading(doc, "Dezco Knoleburn Lore Pack", 0)
    doc.add_paragraph(
        "Organized export of journal entries, Yua story plots, crest/sigil lore, and the party quick reference. "
        "Upload this file to Google Drive → Open with Google Docs. "
        "For clickable tabs, upload Dezco_Knoleburn_Lore_Tabs.xlsx and open with Google Sheets."
    )
    doc.add_paragraph("Suggested Google Docs tabs after upload: Party Ref | Character | Journal 01 | Journal 02 | Journal 03 | Yua Red Black | Yua Mana Well | Crests.")

    sections = [
        ("Party Quick Reference", LORE / "dezco-party-reference.md"),
        ("Character Introduction", LORE / "dezco-knoleburn-character.md"),
        ("Name Meanings", LORE / "knoleburn-names.md"),
        ("Journal 01: After the Charge", LORE / "journal" / "01-after-the-charge.md"),
        ("Journal 02: Leaving Home for Saltrock", LORE / "journal" / "02-leaving-home-for-saltrock.md"),
        ("Journal 03: The Staff Falls", LORE / "journal" / "03-saltrock-staff-falls.md"),
        ("Yua Story Plot: Red and Black", LORE / "yua-story-plots" / "yua-on-red-and-black.md"),
        ("Yua Story Plot: Mana Well", LORE / "yua-story-plots" / "yua-mana-well.md"),
        ("Crests and Sigils (Full Lore)", CRESTS / "README.md"),
    ]

    for title, path in sections:
        doc.add_page_break()
        add_heading(doc, title, 1)
        doc.add_paragraph(f"Source file: {path.relative_to(ROOT)}")
        add_md_ish(doc, read(path))

    # Image appendix paths
    doc.add_page_break()
    add_heading(doc, "Crest and Sigil Image Links", 1)
    doc.add_paragraph("Images live in the repo. After upload to Google Docs, you can Insert → Image from URL using these links:")
    images = sorted(CRESTS.glob("*.png"))
    for img in images:
        url = f"https://github.com/Dralorex/General-Questions/blob/main/assets/crests/alzenhiem/{img.name}"
        raw = f"https://raw.githubusercontent.com/Dralorex/General-Questions/main/assets/crests/alzenhiem/{img.name}"
        doc.add_paragraph(f"{img.name}")
        doc.add_paragraph(f"GitHub: {url}")
        doc.add_paragraph(f"Raw (for insert): {raw}")
        # embed local images into the docx for offline download value
        try:
            doc.add_picture(str(img), width=Inches(3.2))
        except Exception:
            doc.add_paragraph("(image embed skipped)")

    out = OUT / "Dezco_Knoleburn_Lore.docx"
    doc.save(out)
    return out


def main():
    ARTIFACTS.mkdir(parents=True, exist_ok=True)
    xlsx = build_workbook()
    docx_path = build_docx()
    for path in (xlsx, docx_path):
        target = ARTIFACTS / path.name
        target.write_bytes(path.read_bytes())
        print(f"Wrote {path}")
        print(f"Artifact {target}")


if __name__ == "__main__":
    main()
